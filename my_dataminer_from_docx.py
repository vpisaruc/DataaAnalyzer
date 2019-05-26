import docx
import re
from my_database import *

"""Извлечение данных из .docx файлов, преобразование данных в человеческий вид
и запись данных в базу данных файл базы данных analiticdb.dat"""
class DataMiningFromDocx():

    """Преобразование данных всех документов и запись в БД"""
    def _start_data_mining(self):
        cursor.execute("""DELETE FROM tbData;""")
        connection.commit()
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_2.docx', 'Январь', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_3.docx', 'Февраль', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_4.docx', 'Март', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_5.docx', 'Апрель', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_6.docx', 'Май', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_7.docx', 'Июнь', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_8.docx', 'Июль', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_9.docx', 'Август', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_10.docx', 'Сентябрь', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_11.docx', 'Октябрь', '2018')
        self.__from_docx_to_array_2018(r'documents\2018\stat2018_12.docx', 'Ноябрь', '2018')
        self.__from_docx_to_array_2019(r'documents\2019\stat2019_1.docx', 'Январь', '2019')

    """Функция парсит и форматирует данные из ужасного вордовского формата в массив"""
    def __from_docx_to_array_2018(self, path, month, year):
        # Открываю файл
        doc = docx.Document(path)

        # Все таблицы из .docx файла
        tables = doc.tables

        # массив строк таблицы
        string_array = []

        # счетчик ячеек в строке
        cnt = 0

        """Выаскиваю данные по нужной таблице из .docx"""
        for row in tables[5].rows:
            temp = []
            for cell in row.cells:
                cnt += 1
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        # Вычищаю данные от мусора
                        temp.append(re.sub(r'\s+', ' ',
                                    re.sub(r'в (\d),(\d)р\.', '\g<1>\g<2>' + '0',
                                    re.sub(r' в том числе:', '',
                                    re.sub(r' из нее:', '',
                                    re.sub(r' из них:','',re.sub(r'\n','', paragraph.text)))))))
                    if cnt == len(row.cells):
                        string_array.append(temp)
                        cnt = 0

        # Удаление всех пустых элементов
        for i in range(len(string_array)):
            while ' ' in string_array[i]:
                string_array[i].remove(' ')

        # Удаление ненужной строки
        if len(string_array) == 68:
            del string_array[0]
            del string_array[0]
        elif len(string_array) == 67:
            del  string_array[0]

        """Вставляем отредактированные данные в базу данных"""
        for i in range(1, len(string_array)):
            if len(string_array[i]) == 5:
                cursor.execute(" INSERT INTO tbData "
                               "(nameEconomicActivity, sallaryRubles, sallaryLastYearThisMonth, "
                               "sallaryPercentLastMonth, sallaryPercentТationalLevel, curMonth, curYear) "
                               "VALUES (?,?,?,?,?,?,?)",
                               (string_array[i][0], string_array[i][1],
                                 string_array[i][2], string_array[i][3], string_array[i][4], month, year))
                connection.commit()
            elif len(string_array[i]) == 7:
                cursor.execute(" INSERT INTO tbData "
                               "(nameEconomicActivity, sallaryRubles, sallaryLastYearThisMonth, "
                               "sallaryPercentLastMonth, sallaryPercentТationalLevel, "
                               "inRublesJanuaryThisMonth, inPercentJanuaryThisMonthLastYear, curMonth, curYear) "
                               "VALUES (?,?,?,?,?,?,?,?,?)",
                               (string_array[i][0], string_array[i][1],
                                string_array[i][2], string_array[i][3], string_array[i][6],
                                string_array[i][4], string_array[i][5], month, year))
                connection.commit()

    """Функция парсит и форматирует данные из ужасного вордовского формата в массив"""
    def __from_docx_to_array_2019(self, path, month, year):
        # Открываю файл
        doc = docx.Document(path)

        # Все таблицы из .docx файла
        tables = doc.tables

        # массив строк таблицы
        string_array = []

        # счетчик ячеек в строке
        cnt = 0

        """Выаскиваю данные по нужной таблице из .docx"""
        for row in tables[2].rows:
            temp = []
            for cell in row.cells:
                cnt += 1
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        # Вычищаю данные от мусора
                        temp.append(re.sub(r'\s+', ' ',
                                    re.sub(r'в (\d),(\d)р\.', '\g<1>\g<2>' + '0',
                                    re.sub(r' в том числе:', '',
                                    re.sub(r' из нее:', '',
                                           re.sub(r' из них:','',re.sub(r'\n','', paragraph.text)))))))
                    if cnt == len(row.cells):
                        string_array.append(temp)
                        cnt = 0

        # Удаление всех пустых элементов
        for i in range(len(string_array)):
            while ' ' in string_array[i]:
                string_array[i].remove(' ')

        # Удаление ненужной строки
        if len(string_array) == 68:
            del string_array[0]
            del string_array[0]
        elif len(string_array) == 67:
            del string_array[0]

        """Вставляем отредактированные данные в базу данных"""
        for i in range(1, len(string_array)):
            if len(string_array[i]) == 5:
                cursor.execute(" INSERT INTO tbData "
                               "(nameEconomicActivity, sallaryRubles, sallaryLastYearThisMonth, "
                               "sallaryPercentLastMonth, sallaryPercentТationalLevel, curMonth, curYear) "
                               "VALUES (?,?,?,?,?,?,?)",
                               (string_array[i][0], string_array[i][1],
                                 string_array[i][2], string_array[i][3], string_array[i][4], month, year))
                connection.commit()
            elif len(string_array[i]) == 7:
                cursor.execute(" INSERT INTO tbData "
                               "(nameEconomicActivity, sallaryRubles, sallaryLastYearThisMonth, "
                               "sallaryPercentLastMonth, sallaryPercentТationalLevel, "
                               "inRublesJanuaryThisMonth, inPercentJanuaryThisMonthLastYear, curMonth, curYear) "
                               "VALUES (?,?,?,?,?,?,?,?,?)",
                               (string_array[i][0], string_array[i][1],
                                 string_array[i][2], string_array[i][3], string_array[i][6],
                                string_array[i][4], string_array[i][5], month, year))
                connection.commit()



